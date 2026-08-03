#!/usr/bin/env python3
"""Generate .docx and .xlsx downloads for every template, from data/*.yml.

The site is Markdown, but governance artefacts get used in Word (boards, sign-off,
tracked changes) and Excel (registers, logs, scoring). This turns the same data
that renders the pages into files people can actually fill in.

  * every template            -> docs/downloads/<id>.docx
  * templates in spreadsheets.yml -> docs/downloads/<id>.xlsx

Outputs land in docs/downloads/ so MkDocs copies them into the built site and
each page can link to its own file.

Usage:
    python tools/build_downloads.py
    python tools/build_downloads.py --check    # non-zero if anything is missing
"""

from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path

import yaml
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from openpyxl import Workbook
from openpyxl.comments import Comment
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.datavalidation import DataValidation

ROOT = Path(__file__).resolve().parent.parent

# Recommended defaults, loaded in main(); see data/defaults.yml.
DEFAULTS = {}
DATA = ROOT / "data"
OUT = ROOT / "docs" / "downloads"

ACCENT = RGBColor(0x3F, 0x51, 0xB5)
GREY = RGBColor(0x60, 0x60, 0x60)

XL_HEADER_FILL = PatternFill("solid", fgColor="3F51B5")
XL_HEADER_FONT = Font(color="FFFFFF", bold=True, size=10)
XL_INTRO_FONT = Font(italic=True, color="555555", size=10)
XL_BORDER = Border(bottom=Side(style="thin", color="D0D0D0"))


def load(name):
    with open(DATA / name, encoding="utf-8") as fh:
        return yaml.safe_load(fh)


# --------------------------------------------------------------------- Word


def strip_md(text: str) -> str:
    """Markdown inline -> plain text. Word carries its own formatting."""
    # Recommended defaults render as "value (unless: condition)" — the Word
    # copy must carry the recommendation AND when to deviate, or the downloaded
    # document is less useful than the web page it came from.
    def default_sub(m):
        d = DEFAULTS.get(m.group(1))
        if d is None:
            return m.group(0)
        return "%s (unless: %s)" % (
            " ".join(str(d["value"]).split()),
            " ".join(str(d["unless"]).split()),
        )

    text = re.sub(r"\{\{default:([a-z0-9\-]+)\}\}", default_sub, text)
    text = re.sub(r"\{\{([a-z0-9\-]+)\}\}", lambda m: m.group(1).replace("-", " ").title(), text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)      # links
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)            # bold
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)   # italics
    text = re.sub(r"`([^`]+)`", r"\1", text)                  # code
    return text


def add_heading(doc, text, size, bold=True, color=None, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    if color:
        r.font.color.rgb = color
    return p


def add_body(doc, text, italic=False, size=10.5, color=None, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.italic = italic
    if color:
        r.font.color.rgb = color
    return p


def parse_md_table(lines, i):
    """Read a Markdown table starting at lines[i]. Returns (rows, next_index)."""
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        raw = lines[i].strip().strip("|")
        cells = [strip_md(c.strip()) for c in raw.split("|")]
        if not all(re.fullmatch(r":?-{2,}:?", c.replace(" ", "")) for c in cells if c):
            rows.append(cells)
        i += 1
    return rows, i


def write_table(doc, rows):
    if not rows:
        return
    width = max(len(r) for r in rows)
    table = doc.add_table(rows=0, cols=width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        cells = table.add_row().cells
        for ci in range(width):
            text = row[ci] if ci < len(row) else ""
            cell = cells[ci]
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(text)
            run.font.size = Pt(8.5)
            if ri == 0:
                run.font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def render_body_to_docx(doc, body: str):
    """Render the template body Markdown into Word.

    Handles the constructs the templates actually use: headings, paragraphs,
    bullet and numbered lists, tables, checkboxes, and admonitions.
    """
    lines = body.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Admonition: !!! type "Title" followed by an indented block
        m = re.match(r'^!!!\s+(\w+)\s*"?([^"]*)"?\s*$', stripped)
        if m:
            title = m.group(2).strip() or m.group(1).title()
            add_heading(doc, title, 10, bold=True, color=ACCENT, space_before=8)
            i += 1
            buf = []
            while i < len(lines) and (not lines[i].strip() or lines[i].startswith("    ")):
                if lines[i].strip():
                    buf.append(lines[i].strip())
                i += 1
            if buf:
                add_body(doc, strip_md(" ".join(buf)), italic=True, size=10, color=GREY)
            continue

        if stripped.startswith("|"):
            rows, i = parse_md_table(lines, i)
            write_table(doc, rows)
            continue

        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1))
            size = {1: 16, 2: 14, 3: 12, 4: 11}.get(level, 10.5)
            add_heading(doc, strip_md(m.group(2)), size,
                        color=ACCENT if level <= 3 else None)
            i += 1
            continue

        # Checkbox item
        m = re.match(r"^-\s+\[([ xX])\]\s+(.*)$", stripped)
        if m:
            mark = "☒" if m.group(1).lower() == "x" else "☐"
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.space_after = Pt(3)
            r = p.add_run(mark + "  " + strip_md(m.group(2)))
            r.font.size = Pt(10.5)
            i += 1
            continue

        # Bullet
        m = re.match(r"^[-*]\s+(.*)$", stripped)
        if m:
            text = strip_md(m.group(1))
            i += 1
            while i < len(lines) and lines[i].startswith("  ") and lines[i].strip() \
                    and not re.match(r"^\s*[-*|#]|^\s*\d+\.", lines[i].strip()):
                text += " " + strip_md(lines[i].strip())
                i += 1
            p = doc.add_paragraph(text, style="List Bullet")
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                r.font.size = Pt(10.5)
            continue

        # Numbered
        m = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if m:
            text = strip_md(m.group(2))
            i += 1
            while i < len(lines) and lines[i].startswith("   ") and lines[i].strip() \
                    and not re.match(r"^\s*[-*|#]|^\s*\d+\.", lines[i].strip()):
                text += " " + strip_md(lines[i].strip())
                i += 1
            p = doc.add_paragraph(text, style="List Number")
            p.paragraph_format.space_after = Pt(3)
            for r in p.runs:
                r.font.size = Pt(10.5)
            continue

        # Paragraph (join continuation lines)
        buf = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and \
                not re.match(r"^\s*([-*|#>]|\d+\.|!!!)", lines[i].strip()):
            buf.append(lines[i].strip())
            i += 1
        add_body(doc, strip_md(" ".join(buf)))


def build_docx(t: dict, checklist_items=None) -> Document:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Title block
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t["title"])
    r.font.size = Pt(20)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    add_body(doc,
             "GovKit  ·  "
             "Licensed CC BY 4.0  ·  https://lsdeva.github.io/governance-kit/",
             italic=True, size=8.5, color=GREY, space_after=12)

    for label, key in (("Purpose", "purpose"),
                       ("When to use it", "when_to_use"),
                       ("How to use it", "how_to_use")):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(label + ". ")
        run.font.bold = True
        run.font.size = Pt(10.5)
        run2 = p.add_run(strip_md(" ".join(t[key].split())))
        run2.font.size = Pt(10.5)

    doc.add_paragraph()

    if checklist_items is not None:
        render_checklist_docx(doc, checklist_items)
    else:
        render_body_to_docx(doc, t["body"])

    add_heading(doc, "Adaptation notes", 14, color=ACCENT, space_before=14)
    for note in t["adaptation"]:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(note["context"] + ": ")
        r.font.bold = True
        r.font.size = Pt(10.5)
        r2 = p.add_run(strip_md(" ".join(note["note"].split())))
        r2.font.size = Pt(10.5)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(
        "Not legal advice. This template is a head start, not a substitute for "
        "professional judgement. Adapt it to your jurisdiction, sector, and risk "
        "appetite, and have qualified counsel review anything material before you "
        "rely on it."
    )
    r.font.size = Pt(8.5)
    r.font.italic = True
    r.font.color.rgb = GREY
    return doc


def render_checklist_docx(doc, items):
    """The EU AI Act checklist, whose body is generated from questions.yml."""
    add_body(doc,
             "Work through the list and tick what is honestly true today. "
             "Treat every unticked box as a backlog item.",
             space_after=10)
    current = None
    for q in items:
        if q["category"] != current:
            current = q["category"]
            add_heading(doc, current, 12, color=ACCENT, space_before=10)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run("☐  %d. %s" % (q["checklist"], q["text"]))
        r.font.size = Pt(10.5)
        h = doc.add_paragraph()
        h.paragraph_format.left_indent = Inches(0.45)
        h.paragraph_format.space_after = Pt(2)
        hr = h.add_run(" ".join(q["help"].split()))
        hr.font.size = Pt(9)
        hr.font.color.rgb = GREY
        hr.font.italic = True
        if q.get("evidence"):
            e = doc.add_paragraph()
            e.paragraph_format.left_indent = Inches(0.45)
            e.paragraph_format.space_after = Pt(7)
            lbl = e.add_run("Evidence: ")
            lbl.font.size = Pt(9)
            lbl.font.bold = True
            lbl.font.color.rgb = GREY
            ev = e.add_run(" ".join(q["evidence"].split()))
            ev.font.size = Pt(9)
            ev.font.color.rgb = GREY


# -------------------------------------------------------------------- Excel


def col_letter_map(columns):
    return {c["name"]: get_column_letter(i + 1) for i, c in enumerate(columns)}


def write_sheet(ws, spec, title_note=None):
    columns = spec["columns"]
    start = 1

    if title_note:
        ws.cell(row=1, column=1, value=title_note).font = XL_INTRO_FONT
        ws.merge_cells(start_row=1, start_column=1,
                       end_row=1, end_column=max(2, len(columns)))
        ws.row_dimensions[1].height = 28
        ws.cell(row=1, column=1).alignment = Alignment(vertical="center", wrap_text=True)
        start = 3

    header_row = start
    for idx, c in enumerate(columns, start=1):
        cell = ws.cell(row=header_row, column=idx, value=c["name"])
        cell.fill = XL_HEADER_FILL
        cell.font = XL_HEADER_FONT
        cell.alignment = Alignment(vertical="center", wrap_text=True)
        ws.column_dimensions[get_column_letter(idx)].width = c.get("width", 18)
        if c.get("help"):
            cell.comment = Comment(c["help"], "Governance Kit", height=110, width=260)
    ws.row_dimensions[header_row].height = 30

    letters = col_letter_map(columns)
    first_data = header_row + 1

    for r_off, row in enumerate(spec.get("rows", [])):
        r = first_data + r_off
        for c_idx, c in enumerate(columns, start=1):
            value = row[c_idx - 1] if c_idx - 1 < len(row) else ""
            if c.get("formula"):
                value = render_formula(c["formula"], letters, columns, r)
            cell = ws.cell(row=r, column=c_idx, value=value if value != "" else None)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = XL_BORDER
            cell.font = Font(size=10)

    # Blank rows the user fills in, pre-wired with formulas and dropdowns.
    blanks = 30
    for r in range(first_data + len(spec.get("rows", [])),
                   first_data + len(spec.get("rows", [])) + blanks):
        for c_idx, c in enumerate(columns, start=1):
            cell = ws.cell(row=r, column=c_idx)
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = XL_BORDER
            cell.font = Font(size=10)
            if c.get("formula"):
                cell.value = render_formula(c["formula"], letters, columns, r)

    last_row = first_data + len(spec.get("rows", [])) + blanks - 1
    for c_idx, c in enumerate(columns, start=1):
        if not c.get("choices"):
            continue
        dv = DataValidation(
            type="list",
            formula1='"' + ",".join(str(x) for x in c["choices"]) + '"',
            allow_blank=True,
            showDropDown=False,
        )
        dv.error = "Pick a value from the list."
        dv.errorTitle = "Invalid entry"
        ws.add_data_validation(dv)
        letter = get_column_letter(c_idx)
        dv.add("%s%d:%s%d" % (letter, first_data, letter, last_row))

    ws.freeze_panes = ws.cell(row=first_data, column=1)
    ws.auto_filter.ref = "A%d:%s%d" % (header_row, get_column_letter(len(columns)), last_row)
    return ws


def render_formula(template: str, letters, columns, row: int) -> str:
    """Resolve a formula spec into a real Excel formula.

    Formulas reference other columns BY NAME — `{Likelihood}` — never by
    letter. Referring to letters silently breaks the moment a column is
    inserted, and produced a register whose Rating multiplied the wrong two
    cells. Unknown names raise rather than emitting a broken formula.
    """
    def sub(match):
        name = match.group(1)
        if name == "r":
            return str(row)
        if name not in letters:
            raise SystemExit(
                "ERROR: formula references unknown column %r. Known: %s"
                % (name, sorted(letters))
            )
        return letters[name] + str(row)

    return re.sub(r"\{([^{}]+)\}", sub, template)


def build_xlsx(tid: str, spec: dict, template: dict) -> Workbook:
    wb = Workbook()
    ws = wb.active
    ws.title = spec.get("sheet", "Sheet1")[:31]
    write_sheet(ws, spec, title_note=" ".join(spec.get("intro", "").split()) or None)

    for extra in spec.get("extra_sheets", []):
        ws2 = wb.create_sheet(extra["name"][:31])
        write_sheet(ws2, extra)

    info = wb.create_sheet("About")
    info.column_dimensions["A"].width = 100
    rows = [
        (template["title"], True),
        ("", False),
        (" ".join(template["purpose"].split()), False),
        ("", False),
        ("How to use it: " + " ".join(template["how_to_use"].split()), False),
        ("", False),
        ("Columns carry a comment explaining what belongs in them — hover the header.", False),
        ("Grey dropdown columns are constrained to the listed values.", False),
        ("", False),
        ("From GovKit — https://govkit.soa.team/", False),
        ("Licensed CC BY 4.0. Free to use, adapt, and share with attribution.", False),
        ("", False),
        ("Not legal advice. Adapt to your jurisdiction, sector, and risk appetite, "
         "and have qualified counsel review anything material.", False),
    ]
    for i, (text, bold) in enumerate(rows, start=1):
        cell = info.cell(row=i, column=1, value=text)
        cell.font = Font(bold=bold, size=13 if bold else 10)
        cell.alignment = Alignment(wrap_text=True, vertical="top")
    return wb


# --------------------------------------------------------------------- main


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--check", action="store_true",
                    help="verify every expected download exists")
    args = ap.parse_args()

    templates = load("templates.yml")
    questions = load("questions.yml")
    sheets = load("spreadsheets.yml") or {}
    DEFAULTS.update({d["id"]: d for d in (load("defaults.yml") or [])})
    by_id = {t["id"]: t for t in templates}

    unknown = [k for k in sheets if k not in by_id]
    if unknown:
        raise SystemExit("ERROR: spreadsheets.yml references unknown template(s): %s" % unknown)

    checklist = sorted([q for q in questions if q.get("checklist")],
                       key=lambda q: q["checklist"])

    expected = []
    for t in templates:
        expected.append(OUT / (t["id"] + ".docx"))
        if t["id"] in sheets:
            expected.append(OUT / (t["id"] + ".xlsx"))

    if args.check:
        missing = [p for p in expected if not p.exists()]
        if missing:
            print("MISSING downloads:")
            for p in missing:
                print("  %s" % p.relative_to(ROOT))
            print("\nRun: python tools/build_downloads.py")
            return 1
        print("All %d downloads present." % len(expected))
        return 0

    OUT.mkdir(parents=True, exist_ok=True)

    n_doc = n_xls = 0
    for t in templates:
        items = checklist if t.get("generate_body") == "checklist" else None
        build_docx(t, checklist_items=items).save(OUT / (t["id"] + ".docx"))
        n_doc += 1
        if t["id"] in sheets:
            build_xlsx(t["id"], sheets[t["id"]], t).save(OUT / (t["id"] + ".xlsx"))
            n_xls += 1

    print("Generated %d .docx and %d .xlsx in docs/downloads/." % (n_doc, n_xls))
    return 0


if __name__ == "__main__":
    sys.exit(main())
